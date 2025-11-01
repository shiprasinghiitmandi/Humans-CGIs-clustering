import re
import pandas as pd
from docx import Document
import os

# --- File paths ---
input_docx_path = r"/scratch/shipras.sbb.iitmandi/meth_map_codes/chr20highlighted_sequences.docx"
sequence_range_excel_path = r"/scratch/shipras.sbb.iitmandi/sequence_id_ranges.xlsx"
output_docx_path = input_docx_path.replace(".docx", "_renamed.docx")

# --- Step 1: Determine chromosome from docx filename ---
filename = os.path.basename(input_docx_path)
match = re.search(r"chr(\d+|X|Y)", filename)
if not match:
    raise ValueError("Chromosome name not found in filename.")
chromosome = "chr" + match.group(1)

# --- Step 2: Read Excel with chromosome and sequence ID ranges ---
df = pd.read_excel(sequence_range_excel_path)  # <-- uses actual header
df["Chromosome"] = df["Chromosome"].str.strip()

range_row = df[df["Chromosome"] == chromosome]
if range_row.empty:
    raise ValueError(f"No entry found for {chromosome} in Excel file.")

start_id = int(range_row.iloc[0]["Sequence ID Range"].split("-")[0])

# --- Step 3: Open the input DOCX file and rename headers ---
doc = Document(input_docx_path)
current_seq_id = start_id

for para in doc.paragraphs:
    if para.runs and para.runs[0].bold:
        text = para.text.strip()
        if text.startswith(">"):
            para.clear()
            para.add_run(f"sequence {current_seq_id}").bold = True
            current_seq_id += 1

# --- Step 4: Save the modified DOCX file ---
doc.save(output_docx_path)
print(f"Renamed headers saved to: {output_docx_path}")
