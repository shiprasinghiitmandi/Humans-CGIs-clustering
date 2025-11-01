import re
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap

# Input files
text_file_path = r"/scratch/shipras.sbb.iitmandi/meth_map_codes/ATCG_replaced_all/chr20_replaced_ATCG.txt"
excel_file_path = r"/scratch/shipras.sbb.iitmandi/meth_map_codes/excel_cancerlung_meth/chr20_cg_meth.xlsx"
output_docx_path = r"/scratch/shipras.sbb.iitmandi/meth_map_codes/chr20highlighted_sequences.docx"  # Output DOCX file

# Read Excel file, ignoring blank rows
excel_data = pd.read_excel(excel_file_path, header=None).dropna()

# Filter rows where the 4th column (index 3) is 100
filtered_positions = excel_data[excel_data[3] == 100][[0, 1, 2]]  # Columns: Chromosome, Start, End
filtered_positions[[1, 2]] = filtered_positions[[1, 2]].astype(int)  # Ensure integer values

# Read text file
with open(text_file_path, "r") as file:
    lines = file.readlines()

highlight_count = 0  # Counter for highlighted positions

# Process sequences
doc = Document()
current_seq = []
current_header = ""
current_chr = ""
current_range = (0, 0)

for line in lines:
    if line.startswith(">"):  # Sequence header
        if current_seq:  # Process previous sequence
            para = doc.add_paragraph()
            para.add_run(current_header).bold = True
            para.add_run("\n")
            seq_para = doc.add_paragraph()
            for char in current_seq:
                run = seq_para.add_run(char[0] if isinstance(char, tuple) else char)
                if isinstance(char, tuple):  # Highlight background
                    run.font.highlight_color = 6  # Yellow highlight
            para.add_run("\n")
            current_seq = []

        current_header = line.strip()
        match = re.search(r"chr(\d+|X|Y):\s*(\d+)-(\d+)", current_header)
        if match:
            current_chr = "chr" + match.group(1)
            current_range = (int(match.group(2)), int(match.group(3)))
            print(f"Processing sequence: {current_chr} {current_range}")  # Debugging print
    else:
        current_seq.extend(list(line.strip()))  # Store sequence as list

    # Check if any start-end position falls within this sequence range
    for _, (chrom, start, end) in filtered_positions.iterrows():
        if not isinstance(chrom, str):
            continue  # Skip invalid rows
        if chrom.strip() == current_chr and (current_range[0] <= start <= current_range[1]):
            highlight_index = start - current_range[0]  # Convert genome position to sequence index
            if 0 <= highlight_index < len(current_seq):
                print(f"Highlighting position {start} in sequence {current_chr}")  # Debugging print
                current_seq[highlight_index] = (current_seq[highlight_index], True)
                highlight_count += 1

# Save last sequence
if current_seq:
    para = doc.add_paragraph()
    para.add_run(current_header).bold = True
    para.add_run("\n")
    seq_para = doc.add_paragraph()
    for char in current_seq:
        run = seq_para.add_run(char[0] if isinstance(char, tuple) else char)
        if isinstance(char, tuple):  # Highlight background
            run.font.highlight_color = 6  # Yellow highlight
    para.add_run("\n")


doc.save(output_docx_path)


if highlight_count == 0:
    print("No highlights were made.")
else:
    print(f"Total highlighted positions: {highlight_count}")
    print(f"Output saved to {output_docx_path}")
